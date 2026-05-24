from pathlib import Path
from zipfile import ZipFile
import shutil
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/adityagrover/Downloads/profitex 2")
TEMPLATE = Path("/Users/adityagrover/Downloads/FSEReport.docx")
OUT = ROOT / "Profitex_FSE_Report.docx"
WORK = ROOT / "report_assets"
LOGO = WORK / "chitkara_logo.jpeg"


PROJECT = "Profitex"
PROJECT_FULL = "Profitex - Smart Billing & Inventory Management System"
COURSE = "Full Stack Engineering (22CS037)"
GUIDE = "Mohammad Shavez"
AUTHOR = "Lakshay Singla"
DATE = "21.05.26"


def ensure_assets():
    WORK.mkdir(exist_ok=True)
    with ZipFile(TEMPLATE) as zf:
        with zf.open("word/media/image1.jpeg") as src, open(LOGO, "wb") as dst:
            shutil.copyfileobj(src, dst)


def font(size=26, bold=False):
    try:
        return ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", size)
    except Exception:
        return ImageFont.load_default()


def rounded(draw, box, fill, outline=None, radius=18, width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def centered_text(draw, box, text, fnt, fill=(20, 30, 45)):
    lines = []
    for raw in text.split("\n"):
        lines.extend(textwrap.wrap(raw, width=24) or [""])
    heights = []
    widths = []
    for line in lines:
        bb = draw.textbbox((0, 0), line, font=fnt)
        widths.append(bb[2] - bb[0])
        heights.append(bb[3] - bb[1])
    total_h = sum(heights) + (len(lines) - 1) * 6
    y = box[1] + (box[3] - box[1] - total_h) / 2
    for i, line in enumerate(lines):
        x = box[0] + (box[2] - box[0] - widths[i]) / 2
        draw.text((x, y), line, font=fnt, fill=fill)
        y += heights[i] + 6


def save_diagram(name, title, nodes, links=None):
    path = WORK / name
    w, h = 1400, 760
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    d.rectangle((0, 0, w, h), fill=(248, 250, 252))
    d.text((60, 35), title, font=font(42, True), fill=(15, 23, 42))
    box_w, box_h = 285, 105
    for label, x, y, color in nodes:
        rounded(d, (x, y, x + box_w, y + box_h), color, outline=(148, 163, 184), radius=16)
        centered_text(d, (x + 14, y + 12, x + box_w - 14, y + box_h - 12), label, font(25, True))
    if links:
        for a, b in links:
            ax = a[0] + box_w
            ay = a[1] + box_h // 2
            bx = b[0]
            by = b[1] + box_h // 2
            d.line((ax, ay, bx, by), fill=(37, 99, 235), width=5)
            d.polygon([(bx, by), (bx - 18, by - 10), (bx - 18, by + 10)], fill=(37, 99, 235))
    im.save(path)
    return path


def generate_diagrams():
    arch = save_diagram(
        "architecture.png",
        "Profitex System Architecture",
        [
            ("React Client\nPages and Components", 70, 170, (219, 234, 254)),
            ("Axios API Layer\nJWT Token Interceptor", 390, 170, (220, 252, 231)),
            ("Express Server\nProtected REST Routes", 710, 170, (254, 249, 195)),
            ("MongoDB\nMongoose Models", 1030, 170, (255, 228, 230)),
            ("PDFKit Invoice\nDownload Service", 710, 390, (243, 232, 255)),
            ("Uploads Folder\nBusiness Logo Storage", 1030, 390, (224, 242, 254)),
        ],
        [((70, 170), (390, 170)), ((390, 170), (710, 170)), ((710, 170), (1030, 170)), ((710, 390), (1030, 390))],
    )
    er = save_diagram(
        "er_diagram.png",
        "Database Design Overview",
        [
            ("User\nname, email, role", 80, 150, (219, 234, 254)),
            ("Product\nstock, price, GST", 420, 80, (220, 252, 231)),
            ("Invoice\ncustomer, items, total", 420, 250, (254, 249, 195)),
            ("Expense\ntitle, amount, category", 760, 80, (255, 228, 230)),
            ("Purchase\nsupplier, items, cost", 760, 250, (243, 232, 255)),
            ("Company\nGST, address, logo", 1080, 165, (224, 242, 254)),
        ],
        [((80, 150), (420, 80)), ((80, 150), (420, 250)), ((80, 150), (760, 80)), ((80, 150), (760, 250)), ((80, 150), (1080, 165))],
    )
    flow = save_diagram(
        "data_flow.png",
        "Invoice Data Flow",
        [
            ("Login User", 80, 170, (219, 234, 254)),
            ("Select Products\nand Quantity", 385, 170, (220, 252, 231)),
            ("Validate Stock\nand Calculate GST", 690, 170, (254, 249, 195)),
            ("Save Invoice\nUpdate Inventory", 995, 170, (255, 228, 230)),
            ("Download PDF\nView Dashboard", 995, 390, (243, 232, 255)),
        ],
        [((80, 170), (385, 170)), ((385, 170), (690, 170)), ((690, 170), (995, 170))],
    )
    return arch, er, flow


def clear_doc(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def p(doc, text="", size=12, bold=False, italic=False, align=None, space_after=6, color=None):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para


def title(doc, text, size=16):
    return p(doc, text, size=size, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)


def chapter(doc, text):
    p(doc, text, size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)


def sub(doc, text):
    p(doc, text, size=13, bold=True, space_after=6)


def bullet(doc, text):
    para = p(doc, text, size=12, space_after=3)
    para.style = "List Paragraph"
    return para


def table(doc, rows, widths=None):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = t.cell(r, c)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if c == 0 or len(row) <= 3 else WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(str(value))
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            run.bold = r == 0
            if widths:
                cell.width = Inches(widths[c])
    doc.add_paragraph()
    return t


def add_logo(doc):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(LOGO), width=Inches(3.9))


def cover(doc, report_title):
    add_logo(doc)
    title(doc, report_title, 16)
    p(doc, "On", size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    title(doc, PROJECT_FULL, 16)
    p(doc, "Submitted in partial fulfilment of the requirement for the Course", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    title(doc, COURSE, 14)
    p(doc, "of", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    title(doc, "COMPUTER SCIENCE AND ENGINEERING", 14)
    p(doc, "B.E. Batch-2023", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "in", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    title(doc, "MAY-2026", 14)
    table(doc, [["Under the Guidance of", "Submitted By"], [GUIDE, AUTHOR], ["", "B.E. CSE"], ["", "Batch-2023"]], [2.8, 2.8])
    for line in ["DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING", "CHITKARA UNIVERSITY", "PUNJAB"]:
        title(doc, line, 12)
    doc.add_page_break()


def front_matter(doc):
    cover(doc, "INTEGRATED PROJECT REPORT")
    cover(doc, "FULL STACK ENGINEERING PROJECT REPORT")
    title(doc, "CERTIFICATE", 16)
    p(doc, f"This is to be certified that the project entitled \"{PROJECT_FULL}\" has been submitted for the Bachelor of Computer Science Engineering at Chitkara University, Punjab during the academic semester January 2026-May 2026. The project work was carried out under the guidance of the project guide and satisfies the requirements of the course {COURSE}.", size=12, space_after=28)
    p(doc, "Sign. of Project Guide:", size=12, space_after=18)
    p(doc, GUIDE, size=12, bold=True, space_after=30)
    doc.add_page_break()
    title(doc, "CANDIDATE'S DECLARATION", 16)
    p(doc, f"I, {AUTHOR}, B.E.-2023 of Chitkara University, Punjab hereby declare that the Full Stack Engineering project report entitled \"{PROJECT_FULL}\" is an original work carried out for the course {COURSE}. The report has been prepared after analyzing the project source code, database design, frontend implementation, backend APIs, and testing workflow.", size=12, space_after=16)
    table(doc, [["Sign. of Student", "Name", "Batch"], ["", AUTHOR, "B.E. CSE Batch-2023"]], [1.8, 2.2, 2.0])
    p(doc, "Place: Chitkara University", size=12)
    p(doc, f"Date: {DATE}", size=12)
    doc.add_page_break()
    title(doc, "Abstract", 16)
    for para in [
        "Profitex is a full-stack web application developed to support small and medium businesses in managing billing, inventory, purchases, expenses, customer invoices, and financial summaries from one centralized platform.",
        "The system uses a React.js frontend, an Express.js backend, and MongoDB with Mongoose models. It includes user authentication with JWT, protected routes, product CRUD operations, invoice generation, purchase recording, expense categorization, company profile management, and dashboard analytics.",
        "The project addresses common business problems such as manual invoice creation, inaccurate stock records, scattered purchase and expense data, and limited visibility into profit. By connecting sales, purchases, expenses, and inventory, Profitex gives users a clear view of business performance.",
        "A key feature of the project is automated invoice creation. When an invoice is generated, the backend validates stock, calculates GST, updates inventory quantities, stores invoice records, and supports PDF invoice download using PDFKit.",
        "The dashboard summarizes total sales, purchases, expenses, profit, monthly trends, and top-selling products. This makes the system useful not only for transaction entry but also for business decision-making.",
    ]:
        p(doc, para, size=12, space_after=8)
    doc.add_page_break()
    title(doc, "ACKNOWLEDGEMENT", 16)
    for para in [
        "It is our pleasure to be indebted to various people who directly or indirectly contributed to the development of this work and influenced the learning process during the course of study.",
        "We express our sincere gratitude to Chitkara University for providing the opportunity to undergo an Integrated Project as part of the curriculum.",
        f"We are thankful to \"{GUIDE}\" for support, cooperation, and motivation throughout the Full Stack Engineering project work.",
        "We also extend our sincere appreciation to the faculty members of the Department of Computer Science and Engineering for their valuable suggestions and guidance.",
        "Lastly, we would like to thank our parents and friends for their moral support and suggestions that helped improve the quality of this work.",
    ]:
        p(doc, para, size=12, space_after=10)
    doc.add_page_break()
    title(doc, "Table Of Contents", 16)
    rows = [["S.No.", "Contents", "Page No."]]
    contents = [
        ("Introduction", "1-2"),
        ("Objectives", "3"),
        ("Software and Hardware Requirements", "4-5"),
        ("Methods and Design Approach", "6-7"),
        ("Database Analysis, Design and Implementation", "8-10"),
        ("Program Structure Analysis and GUI Construction", "11-14"),
        ("Code Implementation and Database", "15-16"),
        ("System Testing", "17-18"),
        ("Limitation", "19"),
        ("Conclusion", "20"),
        ("Future Scope", "21"),
        ("References/Bibliography", "22"),
    ]
    for i, (item, page_no) in enumerate(contents, 1):
        rows.append([f"{i}.", item, page_no])
    table(doc, rows, [0.8, 4.4, 1.0])
    doc.add_page_break()


def chapters(doc, diagrams):
    arch, er, flow = diagrams
    chapter(doc, "CHAPTER 1 - Introduction")
    sub(doc, "Introduction")
    for para in [
        "Profitex is a Smart Billing and Inventory Management System designed as a full-stack business management platform. The application allows users to register, log in, manage products, create invoices, record purchases, track expenses, search records, configure company information, and view business analytics.",
        "The project is built using the MERN stack pattern. React.js handles the user interface, Express.js provides REST APIs, MongoDB stores application data, and Node.js runs the backend server. Axios is used for client-server communication, JWT is used for authentication, and PDFKit is used for generating downloadable invoice PDFs.",
        "The system is useful for businesses that need a simple digital workflow for sales and inventory. Instead of using separate spreadsheets for products, invoices, purchases, and expenses, Profitex connects these operations in one secured application.",
    ]:
        p(doc, para)
    sub(doc, "Background")
    p(doc, "Small businesses often maintain stock details, purchase history, billing records, GST information, and expense logs manually. Manual record keeping increases the chance of calculation mistakes, stock mismatch, delayed reports, and poor visibility into profit. A full-stack web solution can reduce these issues by automating calculations and storing records in a structured database.")
    sub(doc, "Problem Statement")
    p(doc, "The main problem addressed by Profitex is the absence of an integrated, user-friendly system for managing product inventory, billing, purchase updates, expense records, invoice PDF generation, and dashboard-level business summaries.")
    doc.add_page_break()

    chapter(doc, "CHAPTER 2 - Objectives")
    for item in [
        "To develop a secure web-based business management system for billing and inventory.",
        "To implement user registration and login with token-based authorization.",
        "To maintain product details such as name, category, quantity, price, GST, and supplier.",
        "To generate invoices while validating product stock and calculating GST and grand total.",
        "To automatically reduce product stock after invoice generation and update stock after purchases.",
        "To record expenses by category and generate category-wise and month-wise summaries.",
        "To provide dashboard charts for sales, purchases, expenses, profit, trends, and top-selling products.",
        "To support PDF invoice download and company profile configuration including GST and logo.",
    ]:
        bullet(doc, item)
    doc.add_page_break()

    chapter(doc, "CHAPTER 3 - Software and Hardware Requirements")
    sub(doc, "3.1 Software Requirements")
    table(doc, [
        ["Software/Tool", "Requirement"],
        ["Operating System", "Windows 10/11, macOS, or Linux"],
        ["Frontend Technologies", "React.js, JavaScript, HTML, CSS"],
        ["Backend Technologies", "Node.js, Express.js"],
        ["Database", "MongoDB with Mongoose"],
        ["Libraries", "Axios, Chart.js, React Router, bcryptjs, jsonwebtoken, multer, PDFKit"],
        ["IDE", "Visual Studio Code or similar editor"],
        ["Browser", "Chrome, Edge, Firefox, or Safari"],
    ], [2.2, 4.1])
    sub(doc, "3.2 Hardware Requirements")
    table(doc, [["Hardware", "Specification"], ["Processor", "Intel i5 or higher"], ["RAM", "8 GB recommended"], ["Storage", "256 GB SSD or higher"], ["Internet", "Stable connection for API/database access"]], [2.2, 4.1])
    sub(doc, "3.3 Requirements to run the Application")
    for item in ["Install Node.js and npm.", "Install frontend and backend dependencies using npm install.", "Configure MongoDB connection string in environment variables.", "Add JWT_SECRET and required API configuration in backend environment.", "Start backend server and run the React frontend application.", "Open the application in browser and register/login to access protected modules."]:
        bullet(doc, item)
    doc.add_page_break()

    chapter(doc, "CHAPTER 4 - Methods and Design Approach")
    sub(doc, "4.1 Methodology")
    p(doc, "The project follows an incremental development methodology. Core modules were identified first, then frontend pages, backend routes, database models, authentication middleware, and reporting features were implemented module by module.")
    sub(doc, "Development Steps:")
    for item in ["Requirement Analysis", "UI/UX Design", "Frontend Development", "Backend API Development", "Database Model Design", "Authentication and Route Protection", "Invoice and Inventory Logic", "Dashboard Analytics", "Testing and Debugging"]:
        bullet(doc, item)
    sub(doc, "4.2 System Architecture")
    p(doc, "Profitex uses a three-tier architecture consisting of presentation layer, application layer, and database layer. The React client communicates with Express APIs through Axios. The backend validates JWT tokens, performs business logic, and stores records in MongoDB.")
    doc.add_picture(str(arch), width=Inches(6.2))
    p(doc, "Figure 4.1 System Architecture", size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    chapter(doc, "CHAPTER 5 - Database Analysis, Design and Implementation")
    sub(doc, "5.1 Database Design")
    p(doc, "The database is designed around user-owned records. Each product, invoice, purchase, expense, and company profile is connected to the authenticated user. This prevents records of one user from appearing in another user's workspace.")
    for item in ["Users", "Products", "Invoices", "Purchases", "Expenses", "Company Profiles"]:
        bullet(doc, item)
    sub(doc, "5.2 ER Diagram")
    doc.add_picture(str(er), width=Inches(6.2))
    p(doc, "Figure 5.1 ER Diagram", size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    sub(doc, "5.3 Collection Details")
    table(doc, [
        ["Collection", "Important Fields", "Purpose"],
        ["User", "name, email, password, role", "Stores login identity and role"],
        ["Product", "name, category, quantity, price, gst, supplier", "Maintains inventory records"],
        ["Invoice", "invoiceNumber, customer, items, subtotal, gstTotal, grandTotal", "Stores sales and billing data"],
        ["Purchase", "supplier, items, totalCost", "Stores purchase entries and stock additions"],
        ["Expense", "title, amount, category", "Stores business expenses"],
        ["Company", "name, gstNumber, address, logo", "Stores business profile used in invoices"],
    ], [1.3, 2.7, 2.3])
    doc.add_page_break()

    chapter(doc, "CHAPTER 6 - Program Structure Analysis and GUI Construction")
    sub(doc, "6.1 Program Structure Analysis")
    p(doc, "The project is divided into client and server directories. The client contains React pages, components, context, styling, and API configuration. The server contains routes, controllers, models, database configuration, middleware, and upload storage.")
    
    sub(doc, "6.1.1 System Architecture")
    p(doc, "Profitex follows a three-tier architecture model to ensure clean separation of concerns, high scalability, and secure data access:")
    
    p(doc, "1. Presentation Layer (Frontend)", bold=True)
    p(doc, "Developed using React.js, HTML5, CSS3, JavaScript (ES6+), React Router DOM (client routing), Chart.js (graphs), and Axios (REST client). This layer is responsible for:")
    for item in [
        "User interaction, event handling, and responsive modern dashboard layout.",
        "Secure user authentication client pages (User Registration and Secure Login).",
        "Business Analytics Dashboard (visualizing KPI summaries: Total Sales, Purchases, Expenses, Net Profit, trends, and Top Products).",
        "Product & Inventory Management UI (adding new stock, updating details, viewing quantities, and deleting items).",
        "Billing & Invoicing UI (real-time price calculation, tax/discount adjustments, and PDF invoice generation).",
        "Expense and Purchase tracking interfaces (logging vendor bills and operation overheads).",
        "Real-time keyword search across Products, Invoices, and Expenses.",
        "Company settings page (customizing business details and logo uploads)."
    ]:
        bullet(doc, item)
        
    p(doc, "2. Application Layer (Backend)", bold=True)
    p(doc, "Developed using Node.js and Express.js to construct high-performance RESTful API endpoints. This layer handles:")
    for item in [
        "RESTful API Handling: Routing client HTTP requests to auth, products, invoices, expenses, purchases, and settings controllers.",
        "Authentication Middleware: Validating JWT tokens to protect private routes and authorize sessions.",
        "Request & Business Logic: Calculating billing costs, validating inventory stock counts, and aggregating financial reports.",
        "PDF Generation Service: Compiling branded invoice details into downloadable PDF formats using PDFKit.",
        "Multipart File Upload: Processing and serving corporate logo images using Multer.",
        "System Security: Implementing bcryptjs password hashing and CORS settings to safeguard resources."
    ]:
        bullet(doc, item)
        
    p(doc, "3. Data Layer (Database)", bold=True)
    p(doc, "Developed using MongoDB document database and Mongoose ODM for database connection, schema model definition, and query validation. This layer handles:")
    for item in [
        "Persistent Storage: Securing User, Company, Product, Invoice, Expense, and Purchase collections.",
        "Schema Definition: Enforcing types, default values, and references across models.",
        "Aggregations: Performing pipeline operations for real-time dashboard calculations.",
        "Atomic Updates: Ensuring safe database transitions such as automatic inventory decrement during sales."
    ]:
        bullet(doc, item)

    sub(doc, "6.1.2 Module Structure")
    p(doc, "The system is composed of six distinct business modules designed to maximize reuse and simplify maintenance:")
    
    p(doc, "A. User Authentication & Profile Module", bold=True)
    for item in [
        "Create Account: Register new users with secure password hashing.",
        "Login/Logout: Validate credentials, establish sessions, and manage JWTs.",
        "Profile Guard: Restrict protected actions using authMiddleware validation."
    ]:
        bullet(doc, item)
        
    p(doc, "B. Product & Inventory Management Module", bold=True)
    for item in [
        "Product Catalog: Create, edit, and delete products with categories, GST, and pricing.",
        "Stock Control: Track quantity levels and run real-time regex search filters."
    ]:
        bullet(doc, item)
        
    p(doc, "C. Invoicing & Billing Module", bold=True)
    for item in [
        "Invoice Creation: Compile customer bills with automated GST tax evaluations.",
        "Stock Deduction: Decrease product inventory quantities automatically during a sale.",
        "PDF Receipt Compilation: Compile and export download files via PDFKit."
    ]:
        bullet(doc, item)
        
    p(doc, "D. Expense & Purchase Tracking Module", bold=True)
    for item in [
        "Purchase Logging: Log supplier invoices and automatic stock additions.",
        "Operational Expenses: Track overhead costs categorized by Rent, Salary, Utilities, etc."
    ]:
        bullet(doc, item)
        
    p(doc, "E. Analytics & Dashboard Module", bold=True)
    for item in [
        "Financial KPIs: Tallies Net Profit, Total Sales, Purchases, and Expenses.",
        "Trend Visualizations: Renders responsive monthly line, bar, and pie charts."
    ]:
        bullet(doc, item)
        
    p(doc, "F. Company Profile & Configuration Module", bold=True)
    for item in [
        "Corporate Config: Configures company name, location, contacts, and GSTIN.",
        "Multer Image Store: Uploads business logo for inclusion in billing documents."
    ]:
        bullet(doc, item)

    sub(doc, "6.1.3 Data Flow of a system")
    p(doc, "The transaction and operations data flow in Profitex follows a structured sequence as illustrated below:")
    for item in [
        "1. Authentication: User signs in and retrieves a JWT access token.",
        "2. Request: User triggers an action (e.g. creates an invoice) on the React client UI.",
        "3. Routing: An asynchronous Axios call routes payload data to the server API.",
        "4. Middleware: Express verifies token validity and parses the request payload.",
        "5. Logic: Controller runs stock calculations and computes GST tax summaries.",
        "6. DB Update: Mongoose updates database models and updates product quantities.",
        "7. PDF Generation: Backend renders the PDF invoice template using PDFKit.",
        "8. Response: Server returns JSON success message and updates client dashboard views."
    ]:
        bullet(doc, item)
        
    doc.add_picture(str(flow), width=Inches(6.2))
    p(doc, "Figure 6.1 Invoice Data Flow", size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    sub(doc, "6.2 GUI Construction")
    p(doc, "The user interface is designed using modern styling structures, consistent color branding, and highly responsive elements:")
    
    p(doc, "6.2.1 Authentication Pages", bold=True)
    p(doc, "The Login and Register screens utilize a professional split-view layout featuring local branding elements alongside secured session inputs.")
    
    p(doc, "6.2.2 Business Dashboard Page", bold=True)
    p(doc, "The Business Dashboard Page is the main interactive landing portal of the application for authenticated users.")
    p(doc, "Features:")
    for item in [
        "Navigation bar and responsive sidebar for seamless section routing.",
        "High-level KPI metrics (Total Sales, Total Purchases, Total Expenses, and Net Profit cards).",
        "Monthly trends line chart illustrating sales vs expenses over time.",
        "Comparative budget bar chart and sales-to-expenses ratio pie charts using Chart.js.",
        "Top Selling Products summary tracker displaying quantities sold."
    ]:
        bullet(doc, item)
    p(doc, "Purpose:")
    p(doc, "The page allows business owners to instantly assess their company's financial status, monitor profit margins, visualize expenditure trends, and see high-performing inventory items in a single, responsive layout.")
    
    p(doc, "6.2.3 Product & Inventory Management UI", bold=True)
    p(doc, "Provides lists, filters, stock status categories, and modular input fields to perform CRUD inventory entries.")
    
    p(doc, "6.2.4 Invoices & Billing UI", bold=True)
    p(doc, "Enables real-time customer data entry, product selection dropdowns, dynamic price tallies, and branded PDF invoice downloading.")

    chapter(doc, "CHAPTER 7 - Code Implementation and Database")
    sub(doc, "7.1 Frontend Implementation")
    p(doc, "The frontend is implemented in React.js using React Router for navigation. ProtectedRoute wraps private pages and redirects unauthenticated users to login. The API helper uses Axios and automatically attaches the JWT token from localStorage.")
    sub(doc, "7.2 Backend Implementation")
    p(doc, "The backend uses Express.js with route files for each major module. Controllers contain business logic such as registration, login, invoice creation, purchase recording, expense summary, dashboard calculation, profile update, and company information storage.")
    
    sub(doc, "7.3 API Communication")
    p(doc, "RESTful APIs are used for communication between the frontend client, backend server, database, and upload services:")
    table(doc, [
        ["Method", "API Endpoint", "Function"],
        ["POST", "/api/auth/register", "Register a new user account"],
        ["POST", "/api/auth/login", "Authenticate credentials & retrieve JWT"],
        ["GET", "/api/user/profile", "Retrieve active user profile data"],
        ["PUT", "/api/user/profile", "Modify basic user profile details"],
        ["PUT", "/api/user/password", "Change user account password"],
        ["POST", "/api/products/", "Create a new product in inventory"],
        ["GET", "/api/products/", "Fetch full catalog of inventory items"],
        ["GET", "/api/products/:id", "Fetch specific details of a single product"],
        ["PUT", "/api/products/:id", "Update specifications and quantities of a product"],
        ["DELETE", "/api/products/:id", "Delete product catalog item"],
        ["GET", "/api/products/search/:keyword", "Search inventory records using keyword matching"],
        ["POST", "/api/invoices/", "Generate invoice & deduct product stock"],
        ["GET", "/api/invoices/", "Fetch historical list of customer invoices"],
        ["GET", "/api/invoices/download/:id", "Generate and download branded PDF receipt"],
        ["POST", "/api/purchases/", "Log new product purchase from vendor"],
        ["GET", "/api/purchases/", "Fetch past purchase transaction history"],
        ["POST", "/api/expenses/", "Log operational business expense"],
        ["GET", "/api/expenses/", "Fetch list of logged expenditures"],
        ["GET", "/api/expenses/summary", "Fetch categorized charts data for expenses"],
        ["PUT", "/api/expenses/:id", "Modify an expense record"],
        ["DELETE", "/api/expenses/:id", "Delete an expense entry"],
        ["POST", "/api/company/", "Save company settings & upload logo (Multer)"],
        ["GET", "/api/company/", "Retrieve company details and logo path"],
        ["GET", "/api/dashboard/", "Retrieve dashboard analytics summary"]
    ], [1.1, 2.5, 2.8])
    sub(doc, "7.4 Database Connectivity")
    p(doc, "Database connectivity is configured in server/config/db.js. The backend reads MONGO_URI from environment variables and connects to MongoDB using mongoose.connect. If connection fails, the server logs the error and exits to prevent running without database access.")
    doc.add_page_break()

    chapter(doc, "CHAPTER 8 - System Testing")
    sub(doc, "8.1 Introduction")
    p(doc, "Testing was planned around the main user workflows: authentication, protected page access, product management, invoice generation, purchase recording, expense tracking, dashboard calculation, company profile update, and invoice PDF download.")
    sub(doc, "8.2 Types of Testing Performed")
    
    p(doc, "1. Functional Testing", bold=True)
    p(doc, "Functional testing was performed to verify that all features of the application work correctly.")
    p(doc, "The following functionalities were tested:")
    for item in [
        "User Registration and Secure Login (JWT session validation)",
        "Product Inventory CRUD (add, update, delete, view stock items)",
        "Invoice Generation & Billing Calculations (GST tax, subtotal, grand total)",
        "Inventory Stock Auto-Decrement (on sales/invoicing)",
        "Supplier Purchases & Stock Auto-Increment (on restocking purchases)",
        "Category-wise Expense Tracking & Monthly Summary aggregations",
        "PDF Kit Invoice Compilation and File Downloads",
        "Company profile configurations and logo image uploads (using Multer)"
    ]:
        bullet(doc, item)
    p(doc, "The results confirmed that all modules were functioning according to the project requirements.")

    p(doc, "2. User Interface Testing", bold=True)
    p(doc, "GUI testing was conducted to ensure that the interface is responsive, user-friendly, and visually consistent across different devices.")
    p(doc, "The following aspects were checked:")
    for item in [
        "Interactive Dashboard Charts (Bar, Line, Pie visual representations using Chart.js)",
        "Navigation flow and Protected Route redirections",
        "Inventory search box responsive suggestions",
        "Forms submit validation (e.g. empty fields, phone/email validations, negative quantity blocks)",
        "Responsive layouts across Desktop, Tablet, and Mobile viewport break-points"
    ]:
        bullet(doc, item)
    p(doc, "The interface adapted successfully to desktops, tablets, and mobile devices.")

    p(doc, "3. Database Testing", bold=True)
    p(doc, "Database testing ensured that data was stored, updated, retrieved, and deleted correctly from the database.")
    p(doc, "Testing was performed on collections/tables such as Users, Companies, Products, Invoices, Expenses, and Purchases.")
    p(doc, "The database operations worked successfully without data inconsistency, and relations were properly managed.")

    p(doc, "4. Performance Testing", bold=True)
    p(doc, "Performance testing was conducted to measure the response time and efficiency of the application.")
    p(doc, "The system showed smooth navigation, rapid database query responses (even during complex aggregation pipelines for monthly sales/expense trends), and quick PDF receipt rendering and download response during testing.")
    sub(doc, "8.3 Testing Outcome")
    p(doc, "The testing phase confirmed that Profitex satisfies the major project objectives. The main workflows are modular, protected routes restrict private access, and dashboard values are derived from stored invoice, purchase, and expense records.")
    doc.add_page_break()

    chapter(doc, "CHAPTER 9 - Limitation")
    for item in [
        "The system depends completely on continuous internet connectivity to synchronize with MongoDB Cloud Cluster.",
        "Advanced role-based permissions (e.g. Owner, Manager, Billing Staff) are limited in the current version.",
        "Offline transaction logging and offline invoice storage are not supported.",
        "Security features such as two-factor authentication (2FA) are not implemented in the current version.",
        "Voice-based billing command inputs and voice assistants are not available.",
        "Advanced tax configurations (such as international multi-currency pricing or customized fiscal regions) are limited.",
        "Dedicated native Android and iOS mobile applications are not developed yet.",
        "Automatic invoice sharing via integrated Email and WhatsApp APIs is not fully implemented in the current setup.",
        "Automated low-stock alerts or predictive inventory refilling suggestions are not implemented yet.",
        "Payment gateway integration for real-time customer card/UPI checkout is not supported in the current version."
    ]:
        bullet(doc, item)
    doc.add_page_break()

    chapter(doc, "CHAPTER 10 - Conclusion")
    p(doc, "Profitex successfully implements a smart billing and inventory management platform:")
    for item in [
        "The project helps businesses digitize billing transactions, track expenses, and manage stock easily.",
        "Technologies such as React.js, Node.js, Express.js, MongoDB, Axios, Chart.js, JWT, and PDFKit were successfully integrated.",
        "The system supports major functionalities including secure authentication, cataloging, billing, automated tax calculations, automated inventory stock adjustments, outbound cash logging, and downloadable invoice compilation.",
        "RESTful APIs were implemented for smooth, secure, and authenticated communication between frontend pages and backend endpoints.",
        "The responsive graphical user interface constructed in React and custom CSS improves accessibility across mobile, tablet, and desktop viewports.",
        "Database management was successfully integrated utilizing MongoDB and Mongoose schemas to store and reference multiple relational business entities.",
        "System testing confirmed stable execution, accurate financial calculations, and secure JWT boundaries.",
        "The project improved practical understanding of Full Stack Development, token-based authentication middlewares, dynamic PDF document generation, database schema models, REST APIs, and software engineering concepts.",
        "The project fulfills its primary objective of providing an integrated, simple, and high-performance financial utility for small and medium enterprises.",
        "The modular architecture allows future scalability and enhancements (such as payment gateways, automated notifications, and mobile app developments)."
    ]:
        bullet(doc, item)
    doc.add_page_break()

    chapter(doc, "CHAPTER 11 - Future Scope")
    for item in [
        "Integration of AI models for automated demand forecasting and smart inventory restocking recommendations.",
        "Development of native Android and iOS mobile applications for on-the-go billing and inventory operations.",
        "Addition of multilingual billing and settings support for regional enterprise users.",
        "Implementation of voice-based billing commands (e.g. 'Add item', 'Generate receipt').",
        "Integration of AI chatbot-based customer support and business queries assistance.",
        "Real-time synchronization of regional tax laws, GST rate revisions, and state-level fiscal rules.",
        "Addition of AI-powered financial optimization and expense minimization recommendation systems.",
        "Integration with government taxation portals for direct e-way bill generation and GST return filings.",
        "Cloud-native containerized deployment using Docker and Kubernetes on AWS, Google Cloud, or Microsoft Azure.",
        "Enhancement of security using Two-factor authentication, end-to-end data encryption, and secure cloud backups.",
        "Addition of live billing and accounting expert consultation support.",
        "Implementation of predictive cash flow, net profit, and expense warning analysis.",
        "Integration of email notifications, SMS alerts, and WhatsApp message triggers for sending digital PDF receipts to clients.",
        "Addition of currency conversion and international invoicing features.",
        "Optimization for handling large-scale user traffic and multi-tenant SaaS scaling performance."
    ]:
        bullet(doc, item)
    doc.add_page_break()

    chapter(doc, "CHAPTER 12 - References/Bibliography")
    for item in [
        "React Official Documentation (https://react.dev)",
        "Node.js Official Documentation (https://nodejs.org)",
        "Express.js API Reference (https://expressjs.com)",
        "MongoDB and Mongoose ODM Documentation (https://mongoosejs.com)",
        "Chart.js and React-Chartjs-2 Documentation (https://www.chartjs.org)",
        "JWT (JSON Web Token) Security Specifications (https://jwt.io)",
        "Multer Middleware Upload Guide (https://github.com/expressjs/multer)",
        "PDFKit PDF Generation Library Guide (https://pdfkit.org)",
        "Bcryptjs Password Hashing Standards (https://github.com/dcodeIO/bcrypt.js)",
        "Axios HTTP Client Documentation (https://axios-http.com)"
    ]:
        bullet(doc, item)
    doc.add_page_break()
    p(doc, "The report must consist of following chapters:", size=12, bold=True)
    for item in ["Abstract/Keywords", "Introduction to the project", "Background", "Problem Statement", "Software and Hardware Requirement Specification", "Methods", "Programming/Working Environment", "Requirements to run the application", "Database Analyzing, design and implementation", "Program's Structure Analyzing and GUI Constructing", "Code-Implementation and Database Connections", "System Testing", "Limitations", "Conclusion", "Future Scope", "Bibliography/References"]:
        p(doc, item, size=12, space_after=4)


def main():
    ensure_assets()
    diagrams = generate_diagrams()
    doc = Document(str(TEMPLATE))
    clear_doc(doc)
    set_margins(doc)
    front_matter(doc)
    chapters(doc, diagrams)
    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    main()
